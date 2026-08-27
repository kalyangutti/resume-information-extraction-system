"""
File parser service.

Handles:
- PDF text extraction via PyMuPDF (fitz)
- DOCX text extraction via python-docx
- File type validation
- Empty / corrupted file detection
"""

from __future__ import annotations

import io
from typing import Optional

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class FileParseError(Exception):
    """Raised when a file cannot be parsed."""

    pass


class UnsupportedFileTypeError(FileParseError):
    """Raised when the uploaded file type is not supported."""

    pass


class EmptyFileError(FileParseError):
    """Raised when an uploaded file contains no content."""

    pass


class CorruptedFileError(FileParseError):
    """Raised when a file appears to be corrupted or unreadable."""

    pass


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from an uploaded resume file.

    Args:
        file_bytes: Raw file bytes from the upload.
        filename:   Original filename (used to detect file type).

    Returns:
        Extracted plain text string.

    Raises:
        UnsupportedFileTypeError: If the extension is not .pdf or .docx.
        EmptyFileError:           If the file has no content.
        CorruptedFileError:       If the file cannot be read / parsed.
    """
    if not file_bytes:
        raise EmptyFileError("The uploaded file is empty.")

    extension = _get_extension(filename)

    if extension == ".pdf":
        return _extract_pdf(file_bytes)
    elif extension == ".docx":
        return _extract_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{extension}'. Only PDF and DOCX are accepted."
        )


def validate_file_extension(filename: str) -> None:
    """
    Validate that the filename has a supported extension.

    Args:
        filename: Original filename string.

    Raises:
        UnsupportedFileTypeError: If the extension is not .pdf or .docx.
    """
    ext = _get_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Only .pdf and .docx files are accepted."
        )


def extract_pdf_education_blocks(file_bytes: bytes) -> list[dict]:
    """
    Extract text blocks with positional (x, y) information from a PDF.

    Returns one dict per text block containing:
      - page (int)  : 0-indexed page number
      - x0, y0      : top-left corner coordinates
      - x1, y1      : bottom-right corner coordinates
      - text (str)  : stripped block text

    Only text blocks are returned (image blocks are skipped).
    Returns an empty list if PyMuPDF is unavailable or parsing fails.

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        List of block dicts sorted by (page, y0, x0).
    """
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz  # type: ignore[no-redef]
    except ImportError:
        return []

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception:
        return []

    blocks: list[dict] = []
    for page_num, page in enumerate(doc):
        try:
            raw_blocks = page.get_text("blocks")  # (x0,y0,x1,y1,text,bno,btype)
        except Exception:
            continue
        for b in raw_blocks:
            if b[6] != 0:          # 0 = text block, skip image blocks
                continue
            text = b[4].strip()
            if text:
                blocks.append({
                    "page": page_num,
                    "x0": b[0], "y0": b[1],
                    "x1": b[2], "y1": b[3],
                    "text": text,
                })

    doc.close()
    return sorted(blocks, key=lambda bl: (bl["page"], bl["y0"], bl["x0"]))


def extract_docx_education_tables(file_bytes: bytes) -> list[list[list[str]]]:
    """
    Extract tables from a DOCX file as a nested list structure.

    Returns:
        list[ table ] where each table is a list[ row ] of list[ cell_text ].
        Empty rows and empty tables are excluded.
        Returns an empty list if python-docx is unavailable or parsing fails.

    Args:
        file_bytes: Raw DOCX bytes.
    """
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        return []

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return []

    result: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Deduplicate merged cells (python-docx repeats merged cell text)
            seen: list[str] = []
            for c in cells:
                if not seen or c != seen[-1]:
                    seen.append(c)
            non_empty = [c for c in seen if c]
            if non_empty:
                rows.append(non_empty)
        if rows:
            result.append(rows)

    return result


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _get_extension(filename: str) -> str:
    """Return the lowercase file extension including the dot."""
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF bytes using PyMuPDF.

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        Extracted text.

    Raises:
        CorruptedFileError: If the PDF cannot be opened or is encrypted.
        EmptyFileError:     If the PDF yields no text.
    """
    try:
        try:
            import pymupdf as fitz  # PyMuPDF >= 1.24
        except ImportError:
            import fitz  # PyMuPDF < 1.24 legacy import
    except ImportError as exc:
        raise CorruptedFileError("PyMuPDF is not installed.") from exc

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise CorruptedFileError(f"Could not open PDF file: {exc}") from exc

    if doc.is_encrypted:
        raise CorruptedFileError("The PDF is password-protected and cannot be read.")

    pages_text: list[str] = []
    for page in doc:
        try:
            page_text = page.get_text("text")  # type: ignore[arg-type]
            # Extract PDF hyperlinked URIs (e.g. hyperlinked 'linkedin', 'github' text)
            try:
                links = page.get_links()
                uris = [
                    link["uri"]
                    for link in links
                    if isinstance(link, dict) and "uri" in link and link["uri"]
                ]
                if uris:
                    page_text += "\n" + "\n".join(uris)
            except Exception:
                pass
            pages_text.append(page_text)
        except Exception:
            # Skip pages that cannot be read
            continue

    doc.close()

    text = "\n".join(pages_text)
    if not text.strip():
        raise EmptyFileError(
            "No text could be extracted from the PDF. "
            "It may be a scanned/image-only document."
        )

    return text


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes using python-docx.

    Args:
        file_bytes: Raw DOCX bytes.

    Returns:
        Extracted text.

    Raises:
        CorruptedFileError: If the DOCX file cannot be opened.
        EmptyFileError:     If the DOCX contains no text.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise CorruptedFileError("python-docx is not installed.") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise CorruptedFileError(f"Could not open DOCX file: {exc}") from exc

    paragraphs: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs)
    if not text.strip():
        raise EmptyFileError("No text could be extracted from the DOCX file.")

    return text
